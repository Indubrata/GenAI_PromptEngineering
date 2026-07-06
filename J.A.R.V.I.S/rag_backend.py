import os
import sqlite3
import json
import csv
from datetime import datetime
import numpy as np
import google.generativeai as genai
from dotenv import load_dotenv

# Ensure environment variables are loaded
load_dotenv()

# Configure Gemini API
if os.getenv("GEMINI_API_KEY"):
    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Ensure pypdf and python-docx are imported safely
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

DB_PATH = "rag_knowledge_base.db"

class RAGBackend:
    @staticmethod
    def init_db():
        """Initializes the SQLite database schema if it doesn't exist."""
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        # Documents Table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                filename TEXT PRIMARY KEY,
                file_type TEXT,
                upload_time TEXT,
                total_chunks INTEGER,
                file_size_bytes INTEGER
            )
        """)
        
        # Chunks Table (Storing JSON serialized embeddings)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS chunks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT,
                page_number INTEGER,
                chunk_index INTEGER,
                text TEXT,
                embedding TEXT,
                FOREIGN KEY (filename) REFERENCES documents(filename) ON DELETE CASCADE
            )
        """)
        
        conn.commit()
        conn.close()

    @staticmethod
    def get_embedding(text: str, task_type: str = "retrieval_document") -> list:
        """
        Generates embedding using Gemini's embedding-001 model.
        task_type: 'retrieval_document' for chunk indexing, 'retrieval_query' for query.
        """
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError("GEMINI_API_KEY environment variable is missing.")
        
        try:
            # Task types: 'retrieval_document' or 'retrieval_query'
            result = genai.embed_content(
                model="models/gemini-embedding-001",
                content=text,
                task_type=task_type
            )
            return result["embedding"]
        except Exception as e:
            # Fallback to local mock embeddings or raise error
            raise RuntimeError(f"Gemini Embedding API Error: {str(e)}")

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 800, chunk_overlap: int = 150) -> list:
        """
        Splits text into chunks of chunk_size with chunk_overlap using a sliding window.
        Keeps words intact.
        """
        words = text.split()
        chunks = []
        
        if not words:
            return []
            
        current_word_idx = 0
        while current_word_idx < len(words):
            # Form a chunk
            chunk_words = []
            char_count = 0
            idx = current_word_idx
            
            while idx < len(words) and char_count < chunk_size:
                word = words[idx]
                chunk_words.append(word)
                char_count += len(word) + 1  # Add word length plus space
                idx += 1
                
            chunk_text = " ".join(chunk_words)
            chunks.append(chunk_text)
            
            # Slide window forward: index advances by words corresponding to size - overlap
            # Calculate how many words to advance
            overlap_words_count = 0
            overlap_char_count = 0
            # Backtrack from idx to calculate overlap
            back_idx = idx - 1
            while back_idx >= current_word_idx and overlap_char_count < chunk_overlap:
                word = words[back_idx]
                overlap_char_count += len(word) + 1
                overlap_words_count += 1
                back_idx -= 1
                
            advance = max(1, (idx - current_word_idx) - overlap_words_count)
            current_word_idx += advance
            
        return chunks

    @classmethod
    def parse_pdf(cls, file_bytes) -> list:
        """Parses PDF bytes and returns list of dicts: {'text': str, 'page_number': int}"""
        if PdfReader is None:
            raise ImportError("pypdf is not installed.")
        
        import io
        pdf = PdfReader(io.BytesIO(file_bytes))
        pages_content = []
        
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text and text.strip():
                pages_content.append({
                    "text": text.strip(),
                    "page_number": i + 1
                })
        return pages_content

    @classmethod
    def parse_docx(cls, file_bytes) -> list:
        """Parses DOCX bytes and returns list of paragraphs/tables text."""
        if DocxDocument is None:
            raise ImportError("python-docx is not installed.")
        
        import io
        doc = DocxDocument(io.BytesIO(file_bytes))
        text_blocks = []
        
        # Read paragraphs
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_blocks.append(para.text.strip())
                
        # Read tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_blocks.append(" | ".join(row_text))
                    
        full_text = "\n".join(text_blocks)
        # Word docs usually don't have well-defined digital pages in byte streams, so treat as page 1
        return [{"text": full_text, "page_number": 1}]

    @classmethod
    def parse_csv(cls, file_bytes) -> list:
        """Parses CSV bytes and returns row-by-row text representation."""
        import io
        text_content = io.StringIO(file_bytes.decode('utf-8', errors='ignore'))
        reader = csv.reader(text_content)
        
        rows = []
        headers = []
        for i, row in enumerate(reader):
            if i == 0:
                headers = row
                continue
            row_repr = []
            for col_idx, val in enumerate(row):
                header_name = headers[col_idx] if col_idx < len(headers) else f"Column{col_idx+1}"
                row_repr.append(f"{header_name}: {val}")
            rows.append(f"Row {i}: " + ", ".join(row_repr))
            
        full_text = "\n".join(rows)
        return [{"text": full_text, "page_number": 1}]

    @classmethod
    def parse_txt_md(cls, file_bytes) -> list:
        """Parses plain text / markdown files."""
        text = file_bytes.decode('utf-8', errors='ignore')
        return [{"text": text, "page_number": 1}]

    @classmethod
    def add_document(cls, filename: str, file_bytes: bytes, file_size: int) -> dict:
        """
        Parses document, chunks it, generates embeddings, and indexes it in the SQLite DB.
        """
        cls.init_db()
        
        # Determine file type
        ext = os.path.splitext(filename)[1].lower()
        
        try:
            if ext == '.pdf':
                pages = cls.parse_pdf(file_bytes)
                file_type = "PDF"
            elif ext in ['.docx', '.doc']:
                pages = cls.parse_docx(file_bytes)
                file_type = "DOCX"
            elif ext == '.csv':
                pages = cls.parse_csv(file_bytes)
                file_type = "CSV"
            elif ext in ['.txt', '.md']:
                pages = cls.parse_txt_md(file_bytes)
                file_type = "TEXT/MARKDOWN"
            else:
                raise ValueError(f"Unsupported file format: {ext}")
        except Exception as e:
            return {"success": False, "error": f"Failed to parse document: {str(e)}"}

        if not pages:
            return {"success": False, "error": "Document contains no extractable text."}

        # Step 2: Create chunks keeping track of page numbers
        document_chunks = []
        for page_data in pages:
            page_text = page_data["text"]
            page_num = page_data["page_number"]
            
            chunks = cls.chunk_text(page_text, chunk_size=800, chunk_overlap=150)
            for idx, chunk_text in enumerate(chunks):
                document_chunks.append({
                    "text": chunk_text,
                    "page_number": page_num,
                    "chunk_index": idx
                })
        
        if not document_chunks:
            return {"success": False, "error": "Failed to extract text chunks."}

        # Step 3: Generate embeddings for chunks
        chunk_texts = [c["text"] for c in document_chunks]
        embeddings = []
        
        try:
            # Generate embeddings (we'll process them in batches of 20 to avoid rate limits / token limits)
            batch_size = 20
            for i in range(0, len(chunk_texts), batch_size):
                batch = chunk_texts[i:i+batch_size]
                # Embed batch
                for txt in batch:
                    embeddings.append(cls.get_embedding(txt, task_type="retrieval_document"))
        except Exception as e:
            return {"success": False, "error": f"Failed to generate embeddings: {str(e)}"}

        # Step 4: Save to SQLite
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        try:
            # Delete if exists (overwrite)
            cursor.execute("DELETE FROM documents WHERE filename = ?", (filename,))
            cursor.execute("DELETE FROM chunks WHERE filename = ?", (filename,))
            
            # Insert Document Info
            cursor.execute(
                "INSERT INTO documents (filename, file_type, upload_time, total_chunks, file_size_bytes) VALUES (?, ?, ?, ?, ?)",
                (filename, file_type, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), len(document_chunks), file_size)
            )
            
            # Insert Chunks
            for idx, chunk in enumerate(document_chunks):
                emb_json = json.dumps(embeddings[idx])
                cursor.execute(
                    "INSERT INTO chunks (filename, page_number, chunk_index, text, embedding) VALUES (?, ?, ?, ?, ?)",
                    (filename, chunk["page_number"], chunk["chunk_index"], chunk["text"], emb_json)
                )
                
            conn.commit()
            return {
                "success": True, 
                "filename": filename, 
                "chunks": len(document_chunks), 
                "file_type": file_type
            }
        except Exception as e:
            conn.rollback()
            return {"success": False, "error": f"Database insertion failed: {str(e)}"}
        finally:
            conn.close()

    @classmethod
    def list_documents(cls) -> list:
        """Returns list of all documents currently in the database."""
        cls.init_db()
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        cursor.execute("SELECT filename, file_type, upload_time, total_chunks, file_size_bytes FROM documents ORDER BY upload_time DESC")
        rows = cursor.fetchall()
        
        docs = []
        for r in rows:
            docs.append({
                "filename": r[0],
                "file_type": r[1],
                "upload_time": r[2],
                "total_chunks": r[3],
                "file_size_bytes": r[4]
            })
            
        conn.close()
        return docs

    @classmethod
    def delete_document(cls, filename: str) -> bool:
        """Deletes a document and its chunks from the database."""
        cls.init_db()
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        try:
            cursor.execute("DELETE FROM documents WHERE filename = ?", (filename,))
            cursor.execute("DELETE FROM chunks WHERE filename = ?", (filename,))
            conn.commit()
            return True
        except Exception as e:
            conn.rollback()
            return False
        finally:
            conn.close()

    @classmethod
    def clear_database(cls) -> bool:
        """Clears all documents and chunks from the database."""
        cls.init_db()
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        try:
            cursor.execute("DELETE FROM documents")
            cursor.execute("DELETE FROM chunks")
            conn.commit()
            return True
        except Exception as e:
            conn.rollback()
            return False
        finally:
            conn.close()

    @classmethod
    def cosine_similarity(cls, v1: list, v2: list) -> float:
        """Computes cosine similarity between two numeric lists."""
        a = np.array(v1)
        b = np.array(v2)
        dot = np.dot(a, b)
        norm_a = np.linalg.norm(a)
        norm_b = np.linalg.norm(b)
        
        if norm_a == 0 or norm_b == 0:
            return 0.0
        return float(dot / (norm_a * norm_b))

    @classmethod
    def search(cls, query_text: str, top_k: int = 5, min_similarity: float = 0.0, filter_filenames: list = None) -> list:
        """
        Performs semantic search against stored chunks using Gemini Query Embeddings.
        Returns list of matched chunks with similarity scores.
        """
        cls.init_db()
        
        # 1. Generate query embedding
        try:
            query_emb = cls.get_embedding(query_text, task_type="retrieval_query")
        except Exception as e:
            print(f"Error embedding query: {str(e)}")
            return []
            
        # 2. Fetch chunks from SQLite
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        if filter_filenames:
            # Query with specific filenames
            placeholders = ",".join(["?"] * len(filter_filenames))
            cursor.execute(
                f"SELECT id, filename, page_number, chunk_index, text, embedding FROM chunks WHERE filename IN ({placeholders})",
                tuple(filter_filenames)
            )
        else:
            # Query all chunks
            cursor.execute("SELECT id, filename, page_number, chunk_index, text, embedding FROM chunks")
            
        rows = cursor.fetchall()
        conn.close()
        
        # 3. Compute cosine similarity
        results = []
        for row in rows:
            chunk_id, filename, page_num, chunk_idx, text, emb_json = row
            try:
                emb = json.loads(emb_json)
                similarity = cls.cosine_similarity(query_emb, emb)
                
                if similarity >= min_similarity:
                    results.append({
                        "id": chunk_id,
                        "filename": filename,
                        "page_number": page_num,
                        "chunk_index": chunk_idx,
                        "text": text,
                        "similarity": similarity
                    })
            except Exception as e:
                print(f"Skipping chunk {chunk_id} due to parsing error: {str(e)}")
                
        # 4. Sort and return top_k
        results.sort(key=lambda x: x["similarity"], reverse=True)
        return results[:top_k]
